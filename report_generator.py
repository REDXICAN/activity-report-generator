import pandas as pd
import os
from datetime import datetime
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import Config

class ReportGenerator:
    def __init__(self):
        self.template_path = Config.REPORT_TEMPLATE_PATH
        self.output_path = Config.REPORT_OUTPUT_PATH
    
    def generate_excel_report(self, data, start_date, end_date):
        """Generate Excel report from collected data"""
        
        # Create filename
        filename = self._generate_filename(start_date, end_date, 'xlsx')
        filepath = os.path.join(self.output_path, filename)
        
        # Create Excel writer
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            # Summary sheet
            self._create_summary_sheet(writer, data, start_date, end_date)
            
            # Email activities sheet
            if 'emails' in data:
                self._create_email_sheet(writer, data['emails'])
            
            # GitHub activities sheet
            if 'github' in data:
                self._create_github_sheet(writer, data['github'])
            
            # WhatsApp activities sheet
            if 'whatsapp' in data:
                self._create_whatsapp_sheet(writer, data['whatsapp'])
            
            # Detailed activities sheet
            self._create_detailed_activities_sheet(writer, data)
        
        # Apply formatting
        self._format_excel(filepath)
        
        return filepath
    
    def _create_summary_sheet(self, writer, data, start_date, end_date):
        """Create summary sheet"""
        summary_data = {
            'Periodo': [f"{start_date.strftime('%d/%m/%Y')} - {end_date.strftime('%d/%m/%Y')}"],
            'Total Emails Enviados': [len(data.get('emails', {}).get('sent', []))],
            'Total Emails Recibidos': [len(data.get('emails', {}).get('received', []))],
            'Commits GitHub': [data.get('github', {}).get('stats', {}).get('total_commits', 0)],
            'Pull Requests': [data.get('github', {}).get('stats', {}).get('total_prs', 0)],
            'Clientes Atendidos (WhatsApp)': [data.get('whatsapp', {}).get('unique_customers', 0)],
            'Conversaciones WhatsApp': [len(data.get('whatsapp', {}).get('conversations', []))]
        }
        
        df = pd.DataFrame(summary_data)
        df.to_excel(writer, sheet_name='Resumen', index=False)
    
    def _create_email_sheet(self, writer, email_data):
        """Create email activities sheet"""
        activities = []
        
        # Process sent emails
        for email in email_data.get('sent', []):
            activities.append({
                'Fecha': email.get('datetime', ''),
                'Tipo': 'Enviado',
                'Para': email.get('to', ''),
                'Asunto': email.get('subject', ''),
                'Categoría': self._categorize_email(email)
            })
        
        # Process received emails
        for email in email_data.get('received', []):
            activities.append({
                'Fecha': email.get('datetime', ''),
                'Tipo': 'Recibido',
                'De': email.get('from', ''),
                'Asunto': email.get('subject', ''),
                'Categoría': self._categorize_email(email)
            })
        
        if activities:
            df = pd.DataFrame(activities)
            df.to_excel(writer, sheet_name='Emails', index=False)
    
    def _create_github_sheet(self, writer, github_data):
        """Create GitHub activities sheet"""
        activities = []
        
        # Add commits
        for commit in github_data.get('commits', []):
            activities.append({
                'Fecha': commit['date'],
                'Tipo': 'Commit',
                'Repositorio': commit['repo'],
                'Descripción': commit['message'][:100],
                'Archivos': commit['files_changed'],
                'Adiciones': commit['additions'],
                'Eliminaciones': commit['deletions']
            })
        
        # Add PRs
        for pr in github_data.get('pull_requests', []):
            activities.append({
                'Fecha': pr['created_at'],
                'Tipo': 'Pull Request',
                'Repositorio': pr['repo'],
                'Descripción': pr['title'],
                'Estado': pr['state'],
                'Archivos': '',
                'Adiciones': '',
                'Eliminaciones': ''
            })
        
        if activities:
            df = pd.DataFrame(activities)
            df.to_excel(writer, sheet_name='GitHub', index=False)
    
    def _create_whatsapp_sheet(self, writer, whatsapp_data):
        """Create WhatsApp activities sheet"""
        conversations = []
        
        for conv in whatsapp_data.get('conversations', []):
            conversations.append({
                'Fecha': conv['start_time'],
                'Cliente': conv['customer'],
                'Mensajes': conv['message_count'],
                'Duración': str(conv['end_time'] - conv['start_time']),
                'Temas': ', '.join(conv['topics'])
            })
        
        if conversations:
            df = pd.DataFrame(conversations)
            df.to_excel(writer, sheet_name='WhatsApp', index=False)
    
    def _create_detailed_activities_sheet(self, writer, data):
        """Create detailed activities sheet"""
        all_activities = []
        
        # Add email activities
        for email in data.get('emails', {}).get('sent', []):
            all_activities.append({
                'Fecha': email.get('datetime', ''),
                'Tipo': 'Email',
                'Descripción': f"Email enviado: {email.get('subject', '')}",
                'Detalles': email.get('to', '')
            })
        
        # Add GitHub activities
        for commit in data.get('github', {}).get('commits', []):
            all_activities.append({
                'Fecha': commit['date'],
                'Tipo': 'Desarrollo',
                'Descripción': f"Commit en {commit['repo']}",
                'Detalles': commit['message'][:100]
            })
        
        # Add WhatsApp activities
        for conv in data.get('whatsapp', {}).get('conversations', []):
            all_activities.append({
                'Fecha': conv['start_time'],
                'Tipo': 'Soporte',
                'Descripción': f"Atención cliente WhatsApp",
                'Detalles': f"Cliente: {conv['customer']}, Mensajes: {conv['message_count']}"
            })
        
        if all_activities:
            df = pd.DataFrame(all_activities)
            df.sort_values('Fecha', inplace=True)
            df.to_excel(writer, sheet_name='Actividades Detalladas', index=False)
    
    def _format_excel(self, filepath):
        """Apply formatting to Excel file"""
        wb = load_workbook(filepath)
        
        # Define styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Format each sheet
        for sheet in wb.worksheets:
            # Format headers
            for cell in sheet[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = border
            
            # Auto-adjust column widths
            for column in sheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                
                adjusted_width = min(max_length + 2, 50)
                sheet.column_dimensions[column_letter].width = adjusted_width
            
            # Add borders to all cells with data
            for row in sheet.iter_rows(min_row=2):
                for cell in row:
                    if cell.value:
                        cell.border = border
        
        wb.save(filepath)
    
    def generate_word_report(self, excel_filepath, start_date, end_date):
        """Convert Excel report to Word format"""
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Reporte de Actividades', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add period
        period = doc.add_paragraph()
        period.add_run(f'Período: {start_date.strftime("%d/%m/%Y")} - {end_date.strftime("%d/%m/%Y")}')
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Load Excel data
        excel_data = pd.read_excel(excel_filepath, sheet_name=None)
        
        # Add each sheet as a section
        for sheet_name, df in excel_data.items():
            # Add section heading
            doc.add_heading(sheet_name, 1)
            
            # Add table
            if not df.empty:
                table = doc.add_table(rows=1, cols=len(df.columns))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                header_cells = table.rows[0].cells
                for i, column in enumerate(df.columns):
                    header_cells[i].text = str(column)
                
                # Add data
                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value) if pd.notna(value) else ''
            
            doc.add_page_break()
        
        # Save document
        filename = self._generate_filename(start_date, end_date, 'docx')
        filepath = os.path.join(self.output_path, filename)
        doc.save(filepath)
        
        return filepath
    
    def _generate_filename(self, start_date, end_date, extension):
        """Generate filename for report"""
        # Get next report number
        existing_files = os.listdir(self.output_path)
        numbers = []
        for file in existing_files:
            if file.startswith('Reporte'):
                try:
                    num = int(file.split(' ')[0])
                    numbers.append(num)
                except:
                    pass
        
        next_number = max(numbers) + 1 if numbers else 1
        
        # Format filename
        month_names = {
            1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
            5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
            9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
        }
        
        month = month_names[start_date.month]
        
        if start_date.month == end_date.month:
            filename = f"{next_number} Reporte de Actividades {month} {start_date.day} - {end_date.day}.{extension}"
        else:
            end_month = month_names[end_date.month]
            filename = f"{next_number} Reporte de Actividades {month} {start_date.day} - {end_month} {end_date.day}.{extension}"
        
        return filename
    
    def _categorize_email(self, email):
        """Categorize email based on content"""
        subject = email.get('subject', '').lower()
        body = email.get('body', '').lower()
        
        if any(word in subject + body for word in ['soporte', 'ticket', 'problema']):
            return 'Soporte'
        elif any(word in subject + body for word in ['desarrollo', 'código', 'github']):
            return 'Desarrollo'
        elif any(word in subject + body for word in ['reunión', 'meeting']):
            return 'Reunión'
        elif any(word in subject + body for word in ['reporte', 'informe']):
            return 'Reporte'
        else:
            return 'General'